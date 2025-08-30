import fitz
from docx import Document
from io import BytesIO

class DocumentProcessor:
    def extract_text(self, file_buffer: BytesIO, extension: str) -> str:
        if extension == '.pdf':
            return self._extract_pdf(file_buffer)
        elif extension == '.docx':
            return self._extract_docx(file_buffer)
        elif extension == '.txt':
            return file_buffer.read().decode('utf-8', errors='ignore')
        else:
            raise ValueError(f'Unsupported file type: {extension}')
    
    def _extract_pdf(self, buffer: BytesIO) -> str:
        doc = fitz.open(stream=buffer.read(), filetype='pdf')
        text = ''
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    
    def _extract_docx(self, buffer: BytesIO) -> str:
        doc = Document(buffer)
        text = '\n'.join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                text += ' '.join([cell.text for cell in row.cells]) + '\n'
        return text